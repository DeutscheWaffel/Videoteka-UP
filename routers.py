from datetime import timedelta
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.security import OAuth2PasswordRequestForm
from database import User, Bookmark, CartItem, database, Film, Role
from schemas import UserCreate, UserResponse, Token, UserLogin, AvatarUpdate
from auth import (
    authenticate_user, 
    create_access_token, 
    get_password_hash, 
    get_current_active_user
)
from config import ACCESS_TOKEN_EXPIRE_MINUTES
from fastapi.responses import StreamingResponse
import io
import base64

router = APIRouter()

@router.post("/register", response_model=UserResponse, status_code=status.HTTP_201_CREATED)
async def register(user_data: UserCreate):
    """Регистрация нового пользователя"""
    try:
        # Проверяем, существует ли пользователь с таким username
        existing_user = User.get_or_none(User.username == user_data.username)
        if existing_user:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Пользователь с таким именем уже существует"
            )
        
        # Проверяем, существует ли пользователь с таким email
        existing_email = User.get_or_none(User.email == user_data.email)
        if existing_email:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Пользователь с таким email уже существует"
            )
        
        # Создаем нового пользователя в транзакции
        hashed_password = get_password_hash(user_data.password)
        # Получаем роль "user" по умолчанию
        default_role = Role.get(Role.name == "user")
        
        with database.atomic():
            user = User.create(
                username=user_data.username,
                email=user_data.email,
                hashed_password=hashed_password,
                role=default_role
            )
            # Повторно читаем из БД, чтобы получить значения полей по умолчанию
            user = User.get(User.id == user.id)
        
        return UserResponse.model_validate(user, from_attributes=True)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка при создании пользователя: {str(e)}"
        )

@router.post("/login", response_model=Token)
async def login(user_credentials: UserLogin):
    """Вход в систему"""
    user = authenticate_user(user_credentials.username, user_credentials.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Неверное имя пользователя или пароль",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    
    return {"access_token": access_token, "token_type": "bearer"}

@router.get("/me", response_model=UserResponse)
async def read_users_me(current_user: User = Depends(get_current_active_user)):
    """Получить информацию о текущем пользователе"""
    # Загружаем связанную роль
    user_with_role = User.get(User.id == current_user.id)
    return UserResponse.model_validate(user_with_role, from_attributes=True)

@router.put("/me/avatar", response_model=UserResponse)
async def update_avatar(payload: AvatarUpdate, current_user: User = Depends(get_current_active_user)):
    """Обновить аватар текущего пользователя (base64)"""
    # Небольшая валидация: ограничим размер строки, чтобы не переполнять БД случайно
    if not payload.avatar_base64 or len(payload.avatar_base64) > 5_000_000:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Некорректный размер изображения")
    current_user.avatar_base64 = payload.avatar_base64
    current_user.save()
    return UserResponse.model_validate(current_user, from_attributes=True)

# --- Закладки ---
from pydantic import BaseModel
from typing import List

class BookmarkCreate(BaseModel):
    movie_id: str
    title: str
    author: str | None = None
    price: str | None = None

class BookmarkResponse(BaseModel):
    id: int
    movie_id: str
    title: str
    author: str | None = None
    price: str | None = None

    class Config:
        from_attributes = True

@router.get("/bookmarks", response_model=List[BookmarkResponse])
async def list_bookmarks(current_user: User = Depends(get_current_active_user)):
    items = Bookmark.select().where(Bookmark.user == current_user)
    return [BookmarkResponse.model_validate(item, from_attributes=True) for item in items]

@router.post("/bookmarks", response_model=BookmarkResponse, status_code=status.HTTP_201_CREATED)
async def add_bookmark(payload: BookmarkCreate, current_user: User = Depends(get_current_active_user)):
    try:
        with database.atomic():
            item, created = Bookmark.get_or_create(
                user=current_user,
                movie_id=payload.movie_id,
                defaults={
                    'title': payload.title,
                    'author': payload.author,
                    'price': payload.price,
                }
            )
            if not created:
                # обновим данные, если менялись
                item.title = payload.title
                item.author = payload.author
                item.price = payload.price
                item.save()
        return BookmarkResponse.model_validate(item, from_attributes=True)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Не удалось добавить закладку: {e}")

@router.delete("/bookmarks/{movie_id}", status_code=status.HTTP_204_NO_CONTENT)
async def remove_bookmark(movie_id: str, current_user: User = Depends(get_current_active_user)):
    deleted = Bookmark.delete().where((Bookmark.user == current_user) & (Bookmark.movie_id == movie_id)).execute()
    if deleted == 0:
        raise HTTPException(status_code=404, detail="Закладка не найдена")
    return

# --- Корзина ---
class CartItemCreate(BaseModel):
    movie_id: str
    title: str
    author: str | None = None
    price: str | None = None

class CartItemResponse(BaseModel):
    id: int
    movie_id: str
    title: str
    author: str | None = None
    price: str | None = None

    class Config:
        from_attributes = True

@router.get("/cart", response_model=List[CartItemResponse])
async def list_cart(current_user: User = Depends(get_current_active_user)):
    items = CartItem.select().where(CartItem.user == current_user)
    return [CartItemResponse.model_validate(item, from_attributes=True) for item in items]

@router.post("/cart", response_model=CartItemResponse, status_code=status.HTTP_201_CREATED)
async def add_to_cart(payload: CartItemCreate, current_user: User = Depends(get_current_active_user)):
    try:
        with database.atomic():
            item, created = CartItem.get_or_create(
                user=current_user,
                movie_id=payload.movie_id,
                defaults={
                    'title': payload.title,
                    'author': payload.author,
                    'price': payload.price,
                }
            )
            if not created:
                item.title = payload.title
                item.author = payload.author
                item.price = payload.price
                item.save()
        return CartItemResponse.model_validate(item, from_attributes=True)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Не удалось добавить в корзину: {e}")

@router.delete("/cart/{movie_id}", status_code=status.HTTP_204_NO_CONTENT)
async def remove_from_cart(movie_id: str, current_user: User = Depends(get_current_active_user)):
    deleted = CartItem.delete().where((CartItem.user == current_user) & (CartItem.movie_id == movie_id)).execute()
    if deleted == 0:
        raise HTTPException(status_code=404, detail="Товар не найден в корзине")
    return

# --- Смена пароля ---
from schemas import PasswordChange
from auth import verify_password

@router.put("/me/password", status_code=status.HTTP_204_NO_CONTENT)
async def change_password(payload: PasswordChange, current_user: User = Depends(get_current_active_user)):
    if not verify_password(payload.current_password, current_user.hashed_password):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Текущий пароль неверен")
    if not payload.new_password or len(payload.new_password) < 6:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Новый пароль слишком короткий")
    current_user.hashed_password = get_password_hash(payload.new_password)
    current_user.save()
    return

# --- Фильмы по жанрам ---
class FilmResponse(BaseModel):
    flim_id: int
    title: str
    title_ru: str | None = None
    author: str | None = None
    price: str | None = None
    genre_title: str
    movie_base64: str | None = None

    class Config:
        from_attributes = True

@router.get("/genres/{genre}/films", response_model=List[FilmResponse])
async def get_films_by_genre(genre: str):
    # Приводим жанр к нижнему регистру для соответствия данным
    g = genre.strip().lower()
    q = Film.select().where(Film.genre_title == g)
    return [FilmResponse.model_validate(f, from_attributes=True) for f in q]

@router.get("/films/all", response_model=List[FilmResponse])
async def get_all_films():
    """Получить все фильмы из базы данных"""
    films = Film.select()
    return [FilmResponse.model_validate(f, from_attributes=True) for f in films]

@router.get("/catalog.pdf")
async def download_catalog_pdf():
    """Сформировать PDF каталог со всеми фильмами и постерами"""
    # Импортируем здесь, чтобы не ломать среду при отсутствии зависимости
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Отсутствует зависимость reportlab: {e}")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, title="Каталог фильмов")
    styles = getSampleStyleSheet()

    # Регистрация кириллического шрифта: пробуем проектный и системные шрифты Windows
    from pathlib import Path as _Path
    base_dir = _Path(__file__).resolve().parent
    candidates = [
        base_dir / 'fonts' / 'DejaVuSans.ttf',
        _Path('C:/Windows/Fonts/DejaVuSans.ttf'),
        _Path('C:/Windows/Fonts/arialuni.ttf'),     # Arial Unicode MS
        _Path('C:/Windows/Fonts/seguisym.ttf'),     # Segoe UI Symbol
        _Path('C:/Windows/Fonts/segoeui.ttf'),
        _Path('C:/Windows/Fonts/arial.ttf'),
        _Path('C:/Windows/Fonts/tahoma.ttf'),
    ]
    font_name = None
    for p in candidates:
        try:
            if p.exists():
                name = p.stem  # уникальное имя
                pdfmetrics.registerFont(TTFont(name, str(p)))
                font_name = name
                break
        except Exception:
            continue
    story = []

    films = list(Film.select())

    title_style = styles['Title']
    h_style = styles['Heading2']
    p_style = styles['BodyText']
    if font_name:
        title_style.fontName = font_name
        h_style.fontName = font_name
        p_style.fontName = font_name

    story.append(Paragraph("Каталог фильмов", title_style))
    story.append(Spacer(1, 0.5*cm))

    max_img_w = 10*cm
    max_img_h = 14*cm

    def decode_data_uri(data_uri: str):
        if not data_uri:
            return None
        try:
            # Ожидаем формат data:image/...;base64,XXXX
            if data_uri.startswith('data:') and 'base64,' in data_uri:
                b64 = data_uri.split('base64,', 1)[1]
            else:
                b64 = data_uri
            return base64.b64decode(b64)
        except Exception:
            return None

    def sanitize(text: str) -> str:
        if not isinstance(text, str):
            return text
        # Заменяем символ рубля и иные потенциально проблемные символы
        return text.replace('₽', 'руб.')

    for idx, f in enumerate(films):
        story.append(Paragraph(sanitize(f.title_ru or f.title), h_style))
        meta = []
        if f.author:
            meta.append(f"Режиссёр: {sanitize(f.author)}")
        if f.genre_title:
            meta.append(f"Жанр: {sanitize(f.genre_title)}")
        if f.price:
            meta.append(f"Цена: {sanitize(f.price)}")
        if meta:
            story.append(Paragraph(" | ".join(meta), p_style))
        story.append(Spacer(1, 0.3*cm))

        img_data = decode_data_uri(getattr(f, 'movie_base64', None))
        if img_data:
            try:
                img = Image(io.BytesIO(img_data))
                # Масштабируем под ограничители, сохраняя пропорции
                iw, ih = img.wrap(0, 0)
                scale = min(max_img_w / max(iw, 1), max_img_h / max(ih, 1))
                img.drawWidth = iw * scale
                img.drawHeight = ih * scale
                story.append(img)
                story.append(Spacer(1, 0.5*cm))
            except Exception:
                pass

        # Разделитель между фильмами, новая страница каждые ~4 записи
        story.append(Spacer(1, 0.5*cm))
        if (idx + 1) % 4 == 0 and (idx + 1) != len(films):
            story.append(PageBreak())

    if not films:
        story.append(Paragraph("В каталоге пока нет фильмов", p_style))

    doc.build(story)
    buffer.seek(0)
    return StreamingResponse(buffer, media_type='application/pdf', headers={
        'Content-Disposition': 'attachment; filename="catalog.pdf"'
    })

@router.get("/catalog.docx")
async def download_catalog_docx():
    """Сформировать DOCX каталог со всеми фильмами и постерами"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Отсутствует зависимость python-docx: {e}")

    def decode_data_uri(data_uri: str):
        if not data_uri:
            return None
        try:
            if data_uri.startswith('data:') and 'base64,' in data_uri:
                b64 = data_uri.split('base64,', 1)[1]
            else:
                b64 = data_uri
            return base64.b64decode(b64)
        except Exception:
            return None

    def sanitize(text: str) -> str:
        if not isinstance(text, str):
            return text
        return text.replace('₽', 'руб.')

    films = list(Film.select())

    doc = Document()
    # Заголовок документа
    title = doc.add_heading('Каталог фильмов', level=0)
    # Настроим базовый шрифт на кириллику (если доступен в системе)
    try:
        for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Title']:
            if style_name in [s.name for s in doc.styles]:
                style = doc.styles[style_name]
                if style.font:
                    style.font.name = 'Arial'
                    style.font.size = Pt(11)
    except Exception:
        pass

    for idx, f in enumerate(films):
        doc.add_heading(sanitize(f.title_ru or f.title), level=2)
        meta_parts = []
        if f.author:
            meta_parts.append(f"Режиссёр: {sanitize(f.author)}")
        if f.genre_title:
            meta_parts.append(f"Жанр: {sanitize(f.genre_title)}")
        if f.price:
            meta_parts.append(f"Цена: {sanitize(f.price)}")
        if meta_parts:
            doc.add_paragraph(' | '.join(meta_parts))

        img_data = decode_data_uri(getattr(f, 'movie_base64', None))
        if img_data:
            try:
                # Вставим изображение шириной ~3.5 дюйма
                doc.add_picture(io.BytesIO(img_data), width=Inches(3.5))
            except Exception:
                pass

        # Отступ между записями
        doc.add_paragraph('\n')

    if not films:
        doc.add_paragraph('В каталоге пока нет фильмов')

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return StreamingResponse(out,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': 'attachment; filename="catalog.docx"'}
    )

@router.get("/films/random/{count}", response_model=List[FilmResponse])
async def get_random_films(count: int = 4):
    """Получить случайные фильмы из базы данных"""
    import random
    all_films = list(Film.select())
    random.shuffle(all_films)
    return [FilmResponse.model_validate(f, from_attributes=True) for f in all_films[:count]]

# --- Админ функционал ---
async def get_current_admin_user(current_user: User = Depends(get_current_active_user)) -> User:
    """Получает администратора из текущего пользователя"""
    # Проверяем роль пользователя
    if current_user.role.name != "administrator":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Доступ запрещен. Требуются права администратора"
        )
    return current_user

# ------- Пользователи (админ) -------
from typing import List as _List
from pydantic import BaseModel as _BaseModel

class RoleUpdate(_BaseModel):
    role_name: str

@router.get("/admin/users", response_model=_List[UserResponse])
async def list_users_admin(admin: User = Depends(get_current_admin_user)):
    users = User.select().order_by(User.id)
    return [UserResponse.model_validate(u, from_attributes=True) for u in users]

@router.put("/admin/users/{user_id}/avatar", response_model=UserResponse)
async def admin_update_user_avatar(user_id: int, payload: AvatarUpdate, admin: User = Depends(get_current_admin_user)):
    if not payload.avatar_base64 or len(payload.avatar_base64) > 5_000_000:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Некорректный размер изображения")
    try:
        user = User.get(User.id == user_id)
    except User.DoesNotExist:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    user.avatar_base64 = payload.avatar_base64
    user.save()
    return UserResponse.model_validate(user, from_attributes=True)

@router.put("/admin/users/{user_id}/role", response_model=UserResponse)
async def admin_update_user_role(user_id: int, payload: RoleUpdate, admin: User = Depends(get_current_admin_user)):
    try:
        user = User.get(User.id == user_id)
    except User.DoesNotExist:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    try:
        role = Role.get(Role.name == payload.role_name)
    except Role.DoesNotExist:
        raise HTTPException(status_code=400, detail="Указанная роль не существует")
    user.role = role
    user.save()
    return UserResponse.model_validate(user, from_attributes=True)

class RoleResp(_BaseModel):
    id: int
    name: str
    description: str | None = None

    class Config:
        from_attributes = True

@router.get("/admin/roles", response_model=_List[RoleResp])
async def list_roles_admin(admin: User = Depends(get_current_admin_user)):
    roles = Role.select().order_by(Role.id)
    return [RoleResp.model_validate(r, from_attributes=True) for r in roles]

# Схемы для работы с фильмами
class FilmCreate(BaseModel):
    title: str
    title_ru: str | None = None
    author: str | None = None
    price: str | None = None
    genre_title: str
    movie_base64: str | None = None

@router.post("/admin/films", response_model=FilmResponse, status_code=status.HTTP_201_CREATED)
async def create_film(film_data: FilmCreate, admin: User = Depends(get_current_admin_user)):
    """Создать новый фильм (только для админов)"""
    try:
        film = Film.create(
            title=film_data.title,
            title_ru=film_data.title_ru,
            author=film_data.author,
            price=film_data.price,
            genre_title=film_data.genre_title.lower(),
            movie_base64=film_data.movie_base64
        )
        return FilmResponse.model_validate(film, from_attributes=True)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Не удалось создать фильм: {str(e)}"
        )

@router.delete("/admin/films/{film_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_film(film_id: int, admin: User = Depends(get_current_admin_user)):
    """Удалить фильм (только для админов)"""
    try:
        film = Film.get(Film.flim_id == film_id)
        film.delete_instance()
    except Film.DoesNotExist:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Фильм не найден"
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Не удалось удалить фильм: {str(e)}"
        )
    return
